"""
docx_processor.py — подстановка тегов в docx-шаблон через python-docx.

Поддерживает форматы тегов:
  {{key}}          — простой тег
  {{key:Подсказка}} — тег с подсказкой (подсказка игнорируется при рендере)

Теги могут быть разбиты Word на несколько XML-runs — обрабатывается корректно.
Форматирование первого run тега сохраняется для результирующего текста.
"""

import base64
import io
import os
import re
import copy
from html import unescape as _html_unescape
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches


# Паттерн тега: {{key}} или {{key:подсказка}}
_TAG_RE = re.compile(r"\{\{([^}:]+)(?::[^}]*)?\}\}")

# HTML: <img src="..."> capture group 1 = src value
_IMG_SRC_RE = re.compile(r'<img[^>]+src=["\']([^"\']+)["\'][^>]*/?>',  re.IGNORECASE)
# Strip any remaining HTML tags
_HTML_STRIP_RE = re.compile(r'<[^>]+>')
# Replace <br> with newline before stripping
_BR_RE = re.compile(r'<br\s*/?>', re.IGNORECASE)


def _runs_to_text(runs) -> str:
    return "".join(r.text for r in runs)


def _is_image_value(value: str) -> bool:
    """Проверяет, является ли значение plain base64 data URL картинки."""
    return value.startswith("data:image/")


def _contains_inline_images(value: str) -> bool:
    """Проверяет, содержит ли HTML-значение теги <img>."""
    return "<img" in value


def _strip_html(html: str) -> str:
    """Убрать HTML-теги, заменить <br> на перенос строки."""
    text = _BR_RE.sub("\n", html)
    text = _HTML_STRIP_RE.sub("", text)
    return _html_unescape(text)


def _decode_image(data_url: str) -> tuple[io.BytesIO, str]:
    """Декодирует data URL картинки. Возвращает (BytesIO, mime_type)."""
    header, data = data_url.split(",", 1)
    mime = header.split(";")[0].split(":")[1]  # image/png
    return io.BytesIO(base64.b64decode(data)), mime


def _clear_runs(para) -> None:
    p_elem = para._p
    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)


def _insert_html_with_images(para, html: str) -> None:
    """
    Разбирает HTML со встроенными <img> и вставляет в параграф:
    текстовые части — как runs, картинки — как inline picture runs.
    """
    _clear_runs(para)
    # re.split с capturing group чередует текст и src
    parts = _IMG_SRC_RE.split(html)
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # Текстовая часть
            text = _strip_html(part)
            if text:
                para.add_run(text)
        else:
            # src картинки
            src = part
            if src.startswith("data:image/"):
                try:
                    img_stream, _ = _decode_image(src)
                    run = para.add_run()
                    run.add_picture(img_stream, width=Inches(4))
                except Exception:
                    pass  # пропустить картинку если не удалось декодировать


def _replace_in_paragraph(para, tag_values: dict) -> None:
    """
    Склеивает текст всех runs параграфа, находит теги, заменяет значениями.
    Форматирование берётся от первого run тега.
    Поддерживает значения: plain text, plain data URL, HTML с <img>.
    """
    full_text = _runs_to_text(para.runs)
    if "{{" not in full_text:
        return

    # Ищем первый тег с изображением или HTML-содержимым
    for m in _TAG_RE.finditer(full_text):
        key = m.group(1).strip()
        val = tag_values.get(key, "")
        if not val:
            continue

        if _is_image_value(val):
            # Plain data URL — вставляем как картинку
            _clear_runs(para)
            run = para.add_run()
            img_stream, _ = _decode_image(val)
            try:
                run.add_picture(img_stream, width=Inches(4))
            except Exception:
                run.text = f"[image:{key}]"
            return

        if _contains_inline_images(val):
            # HTML с inline-картинками — вставляем смешанный контент
            _insert_html_with_images(para, val)
            return

    # Обычная замена текста (без картинок)
    def _resolve(m):
        v = tag_values.get(m.group(1).strip(), m.group(0))
        # Если значение содержит HTML (например <br>, <div>), strip тегов
        return _strip_html(v) if "<" in v else v

    new_text = _TAG_RE.sub(_resolve, full_text)
    if new_text == full_text:
        return

    # Сохраним форматирование первого run
    if para.runs:
        first_run = para.runs[0]
        rpr_xml = copy.deepcopy(first_run._r.find(qn("w:rPr")))
    else:
        rpr_xml = None

    # Удаляем все run-элементы из параграфа
    p_elem = para._p
    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)

    # Создаём один новый run с результирующим текстом
    from docx.oxml import OxmlElement
    new_r = OxmlElement("w:r")
    if rpr_xml is not None:
        new_r.append(rpr_xml)
    new_t = OxmlElement("w:t")
    new_t.text = new_text
    if new_text.startswith(" ") or new_text.endswith(" "):
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    p_elem.append(new_r)


def render_docx(template_path: str, output_path: str, tag_values: dict) -> str:
    """
    Заполнить шаблон значениями тегов и сохранить результат.

    Args:
        template_path: путь к исходному .docx шаблону
        output_path:   путь для сохранения заполненного документа
        tag_values:    словарь {key: значение}

    Returns:
        output_path
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document(template_path)

    # Обходим все параграфы документа
    for para in doc.paragraphs:
        _replace_in_paragraph(para, tag_values)

    # Параграфы внутри таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, tag_values)

    doc.save(output_path)
    return output_path
